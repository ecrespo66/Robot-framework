import subprocess
import sys
from docx import *
from docx2pdf import convert
import os
from iBott.files_and_folders.files import File


class Word(File):
    """
    Class to manage Word files it heritages from File class
    Arguments:
        file_path (str): path to the file
    Attributes:
        document (docx.Document): document object
    Methods:
        open(): open the file
        save(path): save the file
        add_heading(text, level): add a heading to the document
        add_paragraph(text, style): add a paragraph to the document
        add_picture(path, size): add a picture to the document
        add_table(matrix): add a table to the document
        read(): read the file and return the text
        convert_to_pdf(path): convert the file to pdf

    """
    def __init__(self, file_path):
        super().__init__(file_path)
        self.document = self.__file()

    def open(self):
        """Open excel file"""

        if os.path.exists(self.path):
            opener = "open" if sys.platform == "darwin" else "xdg-open"
            subprocess.call([opener, self.path])

        elif not os.self.path.exists(self.path):
            self.document.save(self.path)
            opener = "open" if sys.platform == "darwin" else "xdg-open"
            subprocess.call([opener, self.path])

    def save(self, path=None):
        """
        Save word document
        Arguments:
            path {str} -- path to save file (default: {None})
        """

        if path:
            self.document.save(path)
        else:
            self.document.save(self.path)

    def __file(self):
        if os.path.isfile(self.path):
            doc = Document(self.path)
        else:
            doc = Document()
            doc.save(self.path)
        return doc

    def add_heading(self, text, level=0):
        """
        Add Heading to word document
         Arguments:
            text {str} -- text to add
            level {int} -- level of heading (default: {0})
        """
        self.document.add_heading(text, level)
        self.save()

    def add_paragraph(self, text, style=None):
        """
        Add Paragraph to word document
        Arguments:
            text {str} -- text to add
            style {str} -- style of paragraph (default: {None})
        """
        self.document.add_paragraph(text, style=style)
        self.save()

    def add_picture(self, path, size=None):
        """
        Add image to word document
        Arguments:
            path {str} -- path to image
            size {tuple} -- size of image (default: {None})
        """
        if size:
            width = size[0]
            height = size[1]
        else:
            width = None
            height = None
        self.document.add_picture(path, width=width, height=height)
        self.save()

    def add_table(self, matrix):
        """
        Add Table to word document
        Arguments:
            list {list} -- list of lists to add
        """
        table = self.document.add_table(rows=len(matrix), cols=len(matrix[0]))
        for i in range(len(matrix)):
            for j in range(len(matrix[0])):
                table.rows[i].cells[j].text = str(matrix[i][j])
        self.save()

    def read(self):
        """
        Read word document
        Returns:
            list -- list of paragraphs
        """
        data = ""
        fullText = []
        for para in self.document.paragraphs:
            fullText.append(para.text)
            data = '\n'.join(fullText)
        return data

    def convert_to_pdf(self, path=None):
        """
        Convert word document to pdf
        Arguments:
            path {str} -- path to save file (default: {None})
        """

        if path:
            path.replace('.docx', '.pdf')
            convert(path)
        else:
            self.path.replace('.docx', '.pdf')
            convert(self.path)
