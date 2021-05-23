import subprocess
import sys
from docx import *
from docx.shared import Inches
from docx.shared import Cm
from docx2pdf import convert
import os
from .files_activities import File


class Word(File):
    def __init__(self, path):
        """Word document constructor, receives path as parameter"""
        super().__init__(path)
        self.document = self.__file()

    def open(self):
        """Open excel file"""

        if os.path.exists(self.path):
            opener = "open" if sys.platform == "darwin" else "xdg-open"
            subprocess.call([opener, self.path])

        elif not os.self.path.exists(self.path):
            self.document.save(self.path)
            opener = "open" if sys.platform == "darwin" else "xdg-open"
            subprocess.call([opener, self.path])

    def save(self, path=None):
        """Save word document, receives path as optional argument to store document in a different location"""

        if path:
            self.document.save(path)
        else:
            self.document.save(self.path)

    def __file(self):
        if os.path.isfile(self.path):
            doc = Document(self.path)
        else:
            doc = Document()
            doc.save(self.path)
        return doc

    def addHeading(self, text, level=0):
        """Add Heading to word document, receives text and level as optional parameter"""

        self.document.add_heading(text, level)
        self.save()

    def addParagraph(self, text, style=None):
        """Add Paragraph to word document, receives text and style as optional parameter"""

        self.document.add_paragraph(text, style=style)
        self.save()

    def addPicture(self, path, size=None):
        """Add image to word document, receives image path and width as optional parameter"""
        if size:
            width = size[0]
            height = size[1]
        else:
            width = None
            height = None
        self.document.add_picture(path, width=width, height=height)
        self.save()

    def addTable(self, list):
        """Add Table to word document, receives Matrix as parameter"""

        table = self.document.add_table(rows=len(list), cols=len(list[0]))
        for i in range(len(list)):
            for j in range(len(list[0])):
                table.rows[i].cells[j].text = str(list[i][j])
        self.save()

    def read(self):
        """Read word document
        :returns String with output data"""

        data = ""
        fullText = []
        for para in self.document.paragraphs:
            fullText.append(para.text)
            data = '\n'.join(fullText)
        return data

    def convertToPdf(self, path=None):
        """Convert word document to pdf"""

        if path:
            path.replace('.docx', '.pdf')
            convert(path)
        else:
            self.path.replace('.docx', '.pdf')
            convert(self.path)
